from docx import Document
from docx.shared import Inches
import io
import pandas as pd

# Nota: As funções safe_text não são necessárias aqui, pois o docx lida melhor com UTF-8

def criar_docx_relatorio(orcamento_obj, limites, totais_reais, saldo, user_name, frequencia_pagamento) -> bytes:
    """
    Gera o relatório 50-30-20 em formato DOCX (Word).
    Retorna os bytes do documento.
    """
    document = Document()
    
    # --- Título ---
    document.add_heading(f"Gerencie Dindin: Relatorio {orcamento_obj.mes}", 0)
    document.add_paragraph(f"Gerado para: {user_name}")
    document.add_paragraph(f"Frequencia de Pagamento: {frequencia_pagamento}")
    
    document.add_paragraph("\n")

    # --- Resumo Geral e Saldo ---
    document.add_heading("Resumo Geral e Saldo", level=1)
    
    document.add_paragraph(f"Salario Líquido: R$ {orcamento_obj.salario_liquido:,.2f}")
    document.add_paragraph(f"Total Gasto/Alocado: R$ {totais_reais['total_gasto_real']:,.2f}")
    
    # Adicionar saldo final
    saldo_p = document.add_paragraph()
    saldo_p.add_run("SALDO FINAL (Salário - Total Gasto): ").bold = True
    saldo_p.add_run(f"R$ {saldo:,.2f}").bold = True
    
    document.add_paragraph("\n")
    
    # --- Divisão Quinzenal (Se aplicável) ---
    if frequencia_pagamento == 'Quinzenal':
        divisao_quinzenal = orcamento_obj.calcular_divisao_quinzenal(limites) 
        
        document.add_heading("Resumo de Pagamento Quinzenal", level=2)
        document.add_paragraph(f"Salário Líquido Mensal: R$ {orcamento_obj.salario_liquido:,.2f}")
        document.add_paragraph(f"Valor Recebido por Quinzena: R$ {orcamento_obj.salario_liquido / 2:,.2f}")
        
        document.add_paragraph("\n")
        document.add_heading("Limite TOTAL Sugerido para Gastos", level=3)
        
        # Tabela Quinzenal
        table_q = document.add_table(rows=3, cols=3)
        table_q.style = 'Light Shading Accent1'
        hdr_cells = table_q.rows[0].cells
        hdr_cells[0].text = 'Quinzena'
        hdr_cells[1].text = 'Base de Cálculo'
        hdr_cells[2].text = 'Limite Máximo de Gasto'
        
        # 1ª Quinzena
        table_q.rows[1].cells[0].text = '1ª Quinzena'
        table_q.rows[1].cells[1].text = '60% dos Limites Mensais'
        table_q.rows[1].cells[2].text = f"R$ {divisao_quinzenal['Fixas - Início (60%)'] + divisao_quinzenal['Lazer - Início (60%)']:,.2f}"
        
        # 2ª Quinzena
        table_q.rows[2].cells[0].text = '2ª Quinzena'
        table_q.rows[2].cells[1].text = '40% dos Limites Mensais'
        table_q.rows[2].cells[2].text = f"R$ {divisao_quinzenal['Fixas - Meio (40%)'] + divisao_quinzenal['Lazer - Meio (40%)']:,.2f}"
        
        document.add_paragraph("\n")

    # --- Seções 50/30/20 ---
    
    def adicionar_secao_docx(titulo, total_real, limite, despesas, doc):
        doc.add_heading(f"{titulo} (Limite: R$ {limite:,.2f})", level=2)
        
        # Tabela de Despesas
        table_d = doc.add_table(rows=1 + len(despesas), cols=2)
        table_d.style = 'Medium Shading 1'
        
        hdr_cells_d = table_d.rows[0].cells
        hdr_cells_d[0].text = 'Item'
        hdr_cells_d[1].text = 'Valor'
        
        row_idx = 1
        for item, valor in sorted(despesas.items()):
            row_cells = table_d.rows[row_idx].cells
            row_cells[0].text = str(item)
            row_cells[1].text = f"R$ {valor:,.2f}"
            row_idx += 1
            
        doc.add_paragraph(f"TOTAL GASTO: R$ {total_real:,.2f}").bold = True
        
        # Aviso de Limite
        if total_real > limite:
            ultrapassado = total_real - limite
            doc.add_paragraph(f"ATENÇÃO: Você ULTRAPASSOU o limite em R$ {ultrapassado:,.2f}!", style='List Bullet').bold = True
        elif total_real < limite:
            economizado = limite - total_real
            doc.add_paragraph(f"Parabéns! Você economizou R$ {economizado:,.2f} nesta categoria.").italic = True
        
        doc.add_paragraph("\n")

    adicionar_secao_docx("50% Necessidades Fixas (Despesas Fixas)", totais_reais['total_fixas'], limites.get('Necessidades (50%)', 0.0), orcamento_obj.despesas_fixas, document)
    adicionar_secao_docx("30% Desejos e Lazer (Despesas Variáveis)", totais_reais['total_lazer'], limites.get('Desejos/Lazer (30%)', 0.0), orcamento_obj.gastos_lazer, document)
    
    # 20% Poupança
    document.add_heading(f"20% Poupança/Investimento (Meta: R$ {limites.get('Poupança/Investimento (20%)', 0.0):,.2f})", level=2)
    total_poupanca = totais_reais['total_poupanca']
    document.add_paragraph(f"Valor Destinado: R$ {total_poupanca:,.2f}")

    if total_poupanca < limites.get('Poupança/Investimento (20%)', 0.0):
        falta = limites.get('Poupança/Investimento (20%)', 0.0) - total_poupanca
        document.add_paragraph(f"Atenção: Você está R$ {falta:,.2f} abaixo da meta de 20%!", style='List Bullet').bold = True
    
    # --- Saída Final ---
    f = io.BytesIO()
    document.save(f)
    f.seek(0)
    return f.read()

# O código para criar_pdf_relatorio_historico em DOCX seria similar, mas foi omitido por brevidade.
# Lembre-se: Use criar_pdf_relatorio (seu código original corrigido) para o PDF!